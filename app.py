import streamlit as st
import google.generativeai as genai
from docx import Document
import io
from datetime import datetime
import tempfile
import os

# --- ページ設定（美しいUIのベース：wideレイアウトに変更） ---
st.set_page_config(page_title="PDF文字起こし＆Word統合アプリ", layout="wide", initial_sidebar_state="expanded")

# --- カスタムCSS（視認性の向上と不要なUIの完全排除） ---
st.markdown("""
    <style>
    /* 既存の美しいUI設定 */
    .main-header {font-size: 2.2rem; font-weight: bold; color: #1E3A8A; margin-bottom: 0.5rem;}
    .sub-header {font-size: 1.1rem; color: #4B5563; margin-bottom: 2rem;}
    
    /* 開発者用UIの非表示（リスクと無駄の排除） */
    #MainMenu {visibility: hidden;} /* 右上の「…」メニューと英語表記を消す */
    header {visibility: hidden;} /* GitHubアイコンを含む上部ヘッダーを丸ごと消す */
    footer {visibility: hidden;} /* 下部の「Made with Streamlit」を消す */
    
    /* ガイド部分のデザイン微調整 */
    .guide-box {
        background-color: #F3F4F6; 
        padding: 1.5rem; 
        border-radius: 10px; 
        border: 1px solid #E5E7EB;
    }
    </style>
""", unsafe_allow_html=True)

# --- セッション状態の初期化（ログイン状態の保持） ---
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False

# --- パスワード認証画面 ---
if not st.session_state["authenticated"]:
    st.title("🔒 アクセス制限")
    st.write("このアプリを利用するには合言葉を入力してください。")
    
    # Streamlit Secretsから合言葉を取得（クラウド上で後から設定）
    correct_password = st.secrets.get("APP_PASSWORD", "default_password")
    
    password_input = st.text_input("合言葉", type="password")
    
    if st.button("ログイン", type="primary"):
        if password_input == correct_password:
            st.session_state["authenticated"] = True
            st.rerun() # 画面をリロードしてメイン処理へ進む
        else:
            st.error("合言葉が間違っています。")
    
    # 認証されるまではこれ以降のコードを実行しない
    st.stop()

# ==========================================
# これ以降は認証成功時のみ表示・実行される処理
# ==========================================

# --- サイドバー：検索エリア（構造的アップデート） ---
with st.sidebar:
    st.header("🔍 過去データ検索")
    search_query = st.text_input("過去の抽出データを検索（キーワード）")
    if search_query:
        # ※現在はUIのみ。今後の「データ蓄積・データベース連携」フェーズで検索ロジックを実装します
        st.info(f"「{search_query}」の検索結果（※今後のアップデートで実装予定です）")
    
    st.divider()
    st.caption("💡 使い方ガイドは画面右側にも記載しています。")

# SecretsからAPIキーを取得
api_key = st.secrets.get("GEMINI_API_KEY")
if not api_key:
    st.error("システムエラー: APIキーが設定されていません。管理者に連絡してください。")
    st.stop()

genai.configure(api_key=api_key)

# --- メインコンテンツ（ヘッダー） ---
st.markdown('<div class="main-header">📄 PDF文字起こし＆Word統合ツール</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">スキャンした複数のPDFを一度に高精度でテキスト化し、美しいWordファイルに統合します。</div>', unsafe_allow_html=True)

# --- 画面の分割（左側: 75% 作業エリア, 右側: 25% ガイドエリア） ---
main_col, guide_col = st.columns([3, 1])

# ===== 右側：ご利用ガイド =====
with guide_col:
    st.markdown("""
    <div class="guide-box">
        <h4 style="margin-top: 0;">💡 ご利用ガイド</h4>
        <p><b>STEP 1: PDFのアップロード</b><br>
        左側の枠に文字起こししたいPDFをドラッグ＆ドロップします。（複数ファイルの一括選択も可能です）</p>
        <p><b>STEP 2: 統合先Wordの指定（任意）</b><br>
        既存のWordファイルの末尾にテキストを追記したい場合は、右側の枠にアップロードしてください。</p>
        <p><b>STEP 3: 文字起こしの実行</b><br>
        「✨ 文字起こしを開始」ボタンを押すと、AIが全自動でテキストの抽出と、読みやすい形への再構築を行います。</p>
        <p><b>STEP 4: 確認とダウンロード</b><br>
        処理完了後、プレビュー画面が表示されます。内容を確認し、ページ下部のボタンからWordファイルをダウンロードして完了です。</p>
    </div>
    """, unsafe_allow_html=True)
    st.info("📌 ヒント: 文書内の「表」や「段組み」も、AIが自動で構造を判別して綺麗なレイアウトで出力します。")

# ===== 左側：メイン作業エリア =====
with main_col:
    # 構造的アップデートA: 入力エリアの整理（カラム分け）
    with st.container():
        st.write("### 1. ファイルのアップロード")
        upload_col1, upload_col2 = st.columns(2)
        with upload_col1:
            # accept_multiple_files=True に変更し、複数ファイルのリストを受け取る
            uploaded_pdfs = st.file_uploader("📂 PDFファイルをドラッグ＆ドロップ（複数選択可）", type=["pdf"], accept_multiple_files=True)
        with upload_col2:
            uploaded_word = st.file_uploader("📝 統合したい既存のWordファイル（任意）", type=["docx"])
            if uploaded_word:
                st.success(f"統合先ファイル: {uploaded_word.name} の末尾に追記します。")

    # 構造的アップデートB: アクションとプレビューの分離
    st.divider()
    st.write("### 2. 文字起こしの実行")
    # use_container_width=True でボタンを大きく押しやすく
    if st.button("✨ 文字起こしを開始", type="primary", use_container_width=True):
        if not uploaded_pdfs:
            st.error("PDFファイルをアップロードしてください。")
            st.stop()

        # Wordドキュメントの初期化
        if uploaded_word:
            doc = Document(uploaded_word)
            doc.add_page_break() # 末尾に改ページを追加
            doc.add_heading("以下、追加抽出データ", level=1)
        else:
            doc = Document()
            doc.add_heading("文字起こし結果", level=1)

        total_files = len(uploaded_pdfs)
        progress_bar = st.progress(0, text=f"処理を開始します... (0/{total_files})")
        
        all_extracted_texts = [] # プレビュー用のテキストリスト

        try:
            # モデルの初期化（Gemini 2.5 Flashを適用）
            model = genai.GenerativeModel(model_name="gemini-2.5-flash")
            
            # 論理的な再構築と可読性を極限まで高めるプロンプト
            prompt = """
            このPDF文書のテキストを抽出し、人間が最も読みやすい論理的な構造で再構成してください。
            推測や事実の捏造は一切行わず、文書に記載されている情報のみを使用して、以下のルールを厳密に守ること：

            1. 【論理的な再配置】文書全体の論理的な流れを構築すること。文書の「タイトル（題名）」を特定し、その直後に「目次」ブロックを移動させて配置すること。その後ろに本文を順序良く続けること。
            2. 【文章の結合と整形】段組みやページ分割によって途切れた文章は、意味が通るように1つの文章・段落単位で綺麗に結合すること。文中の不自然な改行や、単語間の不要なスペースはすべて削除し、自然で読みやすい日本語の文章に修正すること。
            3. 【ノイズの排除】ヘッダー（例：「〔社会保険通報〕」）やフッター（例：ページ番号）、本文に関係のない記号などはすべて除外すること。
            4. 【表の高度な再現】表（テーブル）が含まれる場合、元の表の行列構造を極限まで正確に読み取り、必ずMarkdown形式の表として出力すること。テキストの羅列に崩さず、極力元の表に近い視覚的構造を維持すること。
            5. 【出力形式】余計な挨拶や前置きは一切出力せず、整形後のテキストデータのみを出力すること。
            """

            # 複数ファイルをループ処理
            for i, uploaded_pdf in enumerate(uploaded_pdfs):
                progress_bar.progress((i) / total_files, text=f"AIが読み取っています... {i+1}/{total_files}件目: {uploaded_pdf.name}")
                
                # PDFを一時ファイルとして保存
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                    tmp_pdf.write(uploaded_pdf.getvalue())
                    tmp_pdf_path = tmp_pdf.name

                # Gemini APIへファイルをアップロード
                sample_file = genai.upload_file(path=tmp_pdf_path, display_name="uploaded_document")
                
                # テキスト抽出の実行
                response = model.generate_content([sample_file, prompt])
                extracted_text = response.text

                # API上のファイルを削除（クリーンアップ）
                genai.delete_file(sample_file.name)
                os.remove(tmp_pdf_path)
                
                # Wordファイルへの書き込み（ファイルごとに見出しと改ページをつける）
                doc.add_heading(f"【ファイル名：{uploaded_pdf.name}】", level=2)
                doc.add_paragraph(extracted_text)
                
                # 最後のファイル以外は改ページを挿入
                if i < total_files - 1:
                    doc.add_page_break()

                # プレビュー用のリストに追加
                all_extracted_texts.append(f"--- 📄 {uploaded_pdf.name} ---\n{extracted_text}\n")

            # 処理完了
            progress_bar.progress(1.0, text=f"処理完了！全 {total_files} 件の文字起こしが終了しました。")
            st.success("すべての文字起こしが完了しました！")
            
            # 構造的アップデートC: 美しいプレビュー（アコーディオン）
            st.markdown("### 📝 抽出結果プレビュー")
            with st.expander("プレビューを確認（クリックで展開）", expanded=True):
                st.text_area("抽出されたテキストデータ:", "\n".join(all_extracted_texts), height=350)

            # メモリ上にWordファイルを保存（ダウンロード用）
            word_io = io.BytesIO()
            doc.save(word_io)
            word_io.seek(0)

            # 構造的アップデートC: ダウンロードへの導線
            st.divider()
            st.subheader("3. データのダウンロード")
            st.write("下のボタンを押すと、お手元のPCやGoogle DriveにWordファイルとして保存できます。")
            
            today_str = datetime.now().strftime("%Y%m%d")
            
            # ファイルが1つの場合と複数の場合で保存名を変える
            if total_files == 1:
                original_name = uploaded_pdfs[0].name.replace(".pdf", "")
                download_filename = f"{today_str}_{original_name}_抽出結果.docx"
            else:
                download_filename = f"{today_str}_複数ファイル一括抽出結果_{total_files}件.docx"

            st.download_button(
                label=f"📥 {download_filename} をダウンロード",
                data=word_io,
                file_name=download_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )

        except Exception as e:
            st.error(f"エラーが発生しました: {e}")
